from __future__ import annotations

import re
from typing import Any

from docx import Document


DEFAULT_GUARDRAIL_NOTICE = (
    "This analysis estimates CV-to-JD fit only and does not guarantee any hiring outcome."
)
MISSING_EVIDENCE_TEXT = "CV evidence was not found in the parsed CV."
MAX_SNIPPET_CHARS = 260
MAX_EVIDENCE_ITEMS = 12

SENSITIVE_KEYS = {
    "access_token",
    "access_token_hash",
    "authorization",
    "bearer",
    "jwt",
    "password",
    "password_hash",
    "bucket",
    "cv_text",
    "file_path",
    "local_path",
    "object_key",
    "raw_cv_text",
    "report_docx_path",
    "s3_key",
    "secret",
    "storage_path",
}

SENSITIVE_PATTERNS = [
    re.compile(r"access_token\s*=\s*[^&\s]+", re.IGNORECASE),
    re.compile(r"access_token=[^&\s]+", re.IGNORECASE),
    re.compile(r"bearer\s+[A-Za-z0-9._\-]+", re.IGNORECASE),
    re.compile(r"jwt\s*[:=]\s*[A-Za-z0-9._\-]+", re.IGNORECASE),
    re.compile(r"https?://[^\s?]+?\?[^\s]+", re.IGNORECASE),
    re.compile(r"(uploads|reports|private)/[^\s]+", re.IGNORECASE),
    re.compile(r"[A-Za-z]:[\\/][^\s]+"),
    re.compile(r"s3://[^\s]+", re.IGNORECASE),
]


def build_docx_report(result: dict, out_path: str):
    safe_result = _scrub_sensitive(result or {})
    evidence_by_id = _evidence_by_id(safe_result)
    include_v3_sections = _has_v3_report_sections(safe_result)

    doc = Document()

    _add_cover(doc, safe_result)
    _add_executive_summary(doc, safe_result)
    _add_score_breakdown(doc, safe_result)
    _add_matched_skills(doc, safe_result, evidence_by_id)
    _add_missing_skills(doc, safe_result)
    _add_evidence_highlights(doc, safe_result)
    _add_improvement_actions(doc, safe_result, include_v3_sections=include_v3_sections)
    if include_v3_sections:
        _add_safe_rewrite_suggestions(doc, safe_result)
        _add_interview_prep(doc, safe_result)
        _add_learning_roadmap(doc, safe_result)
    _add_limitations(doc, safe_result, include_v3_sections=include_v3_sections)
    _add_safe_appendix(doc, safe_result)

    doc.save(out_path)


def _add_cover(doc: Document, result: dict) -> None:
    metadata = result.get("metadata", {}) if isinstance(result.get("metadata"), dict) else {}
    jd_meta = metadata.get("jd", {}) if isinstance(metadata.get("jd"), dict) else {}
    legacy_jd = result.get("jd", {}) if isinstance(result.get("jd"), dict) else {}

    doc.add_heading("AI CV Fit Report", level=1)
    generated_at = metadata.get("generated_at")
    target_role = jd_meta.get("role") or legacy_jd.get("role")
    if generated_at:
        doc.add_paragraph(f"Generated: {_safe_text(generated_at)}")
    if target_role:
        doc.add_paragraph(f"Target role: {_safe_text(target_role)}")


def _add_executive_summary(doc: Document, result: dict) -> None:
    overall = _overall(result)
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(f"Fit score: {_format_score(overall.get('fit_score'))}")
    doc.add_paragraph(f"Fit level: {_safe_text(overall.get('fit_level') or '-')}")
    doc.add_paragraph(f"Summary: {_safe_text(overall.get('summary') or 'Analysis complete.')}")
    doc.add_paragraph(
        "Guardrail: "
        + _safe_text(overall.get("guardrail_notice") or DEFAULT_GUARDRAIL_NOTICE)
    )


def _add_score_breakdown(doc: Document, result: dict) -> None:
    rows = _score_breakdown(result)
    doc.add_heading("Score Breakdown", level=2)
    if not rows:
        doc.add_paragraph("No score breakdown is available.")
        return

    table = doc.add_table(rows=1, cols=4)
    headers = table.rows[0].cells
    headers[0].text = "Component"
    headers[1].text = "Score"
    headers[2].text = "Weight"
    headers[3].text = "Explanation"
    for item in rows:
        cells = table.add_row().cells
        cells[0].text = _safe_text(item.get("label") or item.get("key") or "-")
        cells[1].text = _format_score(item.get("score"))
        cells[2].text = _format_weight(item.get("weight"))
        cells[3].text = _safe_text(item.get("explanation") or "-")


def _add_matched_skills(doc: Document, result: dict, evidence_by_id: dict[str, dict]) -> None:
    matched_skills = _matched_skills(result)
    doc.add_heading("Matched Skills", level=2)
    if not matched_skills:
        doc.add_paragraph("No matched skill groups are available.")
        return

    for item in matched_skills:
        skill = _safe_text(item.get("skill") or "-")
        doc.add_paragraph(
            f"{skill} | requirement={_safe_text(item.get('requirement_type') or '-')} | "
            f"confidence={_format_score(item.get('confidence'))}"
        )
        doc.add_paragraph(f"JD requirement: {_safe_text(item.get('jd_requirement') or '-')}")
        cv_ids = item.get("cv_evidence_ids") if isinstance(item.get("cv_evidence_ids"), list) else []
        snippets = [_evidence_snippet(evidence_by_id[item_id]) for item_id in cv_ids if item_id in evidence_by_id]
        if snippets:
            for snippet in snippets[:3]:
                doc.add_paragraph(f"CV evidence: {snippet}")
        else:
            doc.add_paragraph(MISSING_EVIDENCE_TEXT)


def _add_missing_skills(doc: Document, result: dict) -> None:
    missing_skills = _missing_skills(result)
    doc.add_heading("Missing Skills / Gaps", level=2)
    if not missing_skills:
        doc.add_paragraph("No major missing skill evidence was detected.")
        return

    for item in missing_skills:
        doc.add_paragraph(
            f"{_safe_text(item.get('skill') or '-')} | "
            f"requirement={_safe_text(item.get('requirement_type') or '-')} | "
            f"severity={_safe_text(item.get('severity') or '-')}"
        )
        doc.add_paragraph(f"JD requirement: {_safe_text(item.get('jd_requirement') or '-')}")
        doc.add_paragraph(f"Reason: {_safe_text(item.get('reason') or MISSING_EVIDENCE_TEXT)}")
        suggestion = item.get("suggestion")
        if suggestion:
            doc.add_paragraph(f"Suggestion: {_safe_text(suggestion)}")


def _add_evidence_highlights(doc: Document, result: dict) -> None:
    evidence = _evidence(result)
    doc.add_heading("Evidence Highlights", level=2)
    if not evidence:
        doc.add_paragraph("No evidence snippets are available.")
        return

    selected = _selected_evidence(evidence)
    for item in selected[:MAX_EVIDENCE_ITEMS]:
        source = _safe_text(item.get("source") or "-")
        kind = _safe_text(item.get("kind") or item.get("type") or "-")
        doc.add_paragraph(f"{source.upper()} {kind}: {_evidence_snippet(item)}")
        if item.get("jd_requirement") and source == "cv":
            doc.add_paragraph(f"JD requirement: {_safe_text(item.get('jd_requirement'))}")
        if item.get("similarity") is not None:
            doc.add_paragraph(f"Similarity: {_format_score(item.get('similarity'))}")


def _add_improvement_actions(doc: Document, result: dict, *, include_v3_sections: bool) -> None:
    actions = _improvement_actions(result)
    if include_v3_sections:
        doc.add_heading("Improvement Action Plan", level=2)
        doc.add_paragraph("Improvement Actions")
    else:
        doc.add_heading("Improvement Actions", level=2)
    if not actions:
        doc.add_paragraph("No improvement actions are available.")
        return

    for item in actions:
        doc.add_paragraph(
            f"{_safe_text(item.get('priority') or '-').upper()}: {_safe_text(item.get('title') or '-')}"
        )
        linked_skill = item.get("linked_skill") or item.get("related_skill")
        if linked_skill:
            doc.add_paragraph(f"Linked skill: {_safe_text(linked_skill)}")
        if item.get("reason"):
            doc.add_paragraph(f"Reason: {_safe_text(item.get('reason'))}")
        suggestion = item.get("safe_suggestion") or item.get("suggestion")
        doc.add_paragraph(f"Safe suggestion: {_safe_text(suggestion or '-')}")
        doc.add_paragraph(f"Status: {_safe_text(item.get('status') or '-')}")
        guardrail = item.get("guardrail") or "Only add this if it is true. Do not fabricate skills or experience."
        if item.get("do_not_fabricate") is True and "fabricate" not in str(guardrail).lower():
            guardrail = f"{guardrail} Do not fabricate skills or experience."
        doc.add_paragraph(f"Guardrail: {_safe_text(guardrail)}")


def _add_safe_rewrite_suggestions(doc: Document, result: dict) -> None:
    suggestions = _safe_rewrite_suggestions(result)
    doc.add_heading("Safe Rewrite Suggestions", level=2)
    if not suggestions:
        doc.add_paragraph("No safe rewrite suggestions are available.")
        return

    for index, item in enumerate(suggestions, start=1):
        doc.add_paragraph(f"Suggestion {index}")
        source_evidence = _safe_list(item.get("source_evidence"))
        if source_evidence:
            doc.add_paragraph(
                "Source evidence: "
                + "; ".join(_truncate(_safe_text(value), MAX_SNIPPET_CHARS) for value in source_evidence[:3])
            )
        doc.add_paragraph(f"Suggested structure: {_safe_text(item.get('suggested_structure') or '-')}")
        missing_context = _safe_list(item.get("missing_context_to_confirm"))
        if missing_context:
            doc.add_paragraph("Missing context to confirm: " + "; ".join(_safe_text(value) for value in missing_context[:6]))
        warning = item.get("warning") or "Only use details that are true and can be defended in an interview."
        doc.add_paragraph(f"Warning: {_safe_text(warning)}")
        if item.get("do_not_fabricate") is True:
            doc.add_paragraph("Do not fabricate skills, experience, metrics, employers, dates, or certifications.")


def _add_interview_prep(doc: Document, result: dict) -> None:
    questions = _interview_prep(result)
    doc.add_heading("Interview Prep", level=2)
    if not questions:
        doc.add_paragraph("No interview prep questions are available.")
        return

    for item in questions:
        doc.add_paragraph(f"Question: {_safe_text(item.get('question') or '-')}")
        doc.add_paragraph(f"Type: {_safe_text(item.get('type') or '-')}")
        doc.add_paragraph(f"Why this question: {_safe_text(item.get('why_this_question') or '-')}")
        doc.add_paragraph(f"Related JD requirement: {_safe_text(item.get('related_jd_requirement') or '-')}")
        related_cv = _safe_list(item.get("related_cv_evidence"))
        if related_cv:
            doc.add_paragraph("Related CV evidence: " + "; ".join(_safe_text(value) for value in related_cv[:4]))
        else:
            doc.add_paragraph(MISSING_EVIDENCE_TEXT)
        outline = _safe_list(item.get("suggested_answer_outline"))
        if outline:
            doc.add_paragraph("Suggested answer outline:")
            for point in outline[:6]:
                doc.add_paragraph(_safe_text(point), style="List Bullet")
        doc.add_paragraph(f"Risk if user cannot answer: {_safe_text(item.get('risk_if_user_cannot_answer') or '-')}")


def _add_learning_roadmap(doc: Document, result: dict) -> None:
    items = _learning_roadmap(result)
    doc.add_heading("Learning Roadmap", level=2)
    if not items:
        doc.add_paragraph("No learning roadmap items are available.")
        return

    for item in items:
        doc.add_paragraph(
            f"{_safe_text(item.get('priority') or '-').upper()}: {_safe_text(item.get('skill') or '-')}"
        )
        doc.add_paragraph(f"Why: {_safe_text(item.get('why') or '-')}")
        topics = _safe_list(item.get("topics"))
        if topics:
            doc.add_paragraph("Topics: " + "; ".join(_safe_text(value) for value in topics[:8]))
        doc.add_paragraph(f"Mini project: {_safe_text(item.get('mini_project') or '-')}")
        doc.add_paragraph(f"Estimated effort: {_safe_text(item.get('estimated_effort') or '-')}")
        doc.add_paragraph(
            "CV evidence to add after learning: "
            + _safe_text(item.get("cv_evidence_to_add_after_learning") or "-")
        )
        if item.get("do_not_claim_until_completed") is True:
            doc.add_paragraph("Do not claim this skill until you have completed real work you can explain.")


def _add_limitations(doc: Document, result: dict, *, include_v3_sections: bool) -> None:
    if include_v3_sections:
        doc.add_heading("Limitations / Safety Notes", level=2)
        doc.add_paragraph("Limitations")
    else:
        doc.add_heading("Limitations", level=2)
    for item in _limitations(result):
        doc.add_paragraph(_safe_text(item), style=None)


def _add_safe_appendix(doc: Document, result: dict) -> None:
    metadata = result.get("metadata", {}) if isinstance(result.get("metadata"), dict) else {}
    cv_meta = metadata.get("cv", {}) if isinstance(metadata.get("cv"), dict) else {}
    doc.add_heading("Appendix: Safe Metadata", level=2)
    doc.add_paragraph(f"schema_version: {_safe_text(result.get('schema_version') or '-')}")
    doc.add_paragraph(f"scorer_version: {_safe_text(metadata.get('scorer_version') or '-')}")
    doc.add_paragraph(f"parser_confidence: {_format_score(cv_meta.get('parsed_confidence'))}")
    detected_skills = cv_meta.get("skills_detected", [])
    count = len(detected_skills) if isinstance(detected_skills, list) else 0
    doc.add_paragraph(f"detected_skills_count: {count}")


def _overall(result: dict) -> dict:
    overall = result.get("overall")
    if isinstance(overall, dict):
        return overall
    fit_score = _extract_fit_score(result)
    return {
        "fit_score": fit_score,
        "fit_level": _fit_level(fit_score),
        "summary": "Analysis complete.",
        "guardrail_notice": DEFAULT_GUARDRAIL_NOTICE,
    }


def _score_breakdown(result: dict) -> list[dict]:
    breakdown = result.get("score_breakdown")
    if isinstance(breakdown, list) and breakdown:
        return [_scrub_sensitive(item) for item in breakdown if isinstance(item, dict)]

    scores = result.get("scores", {}) if isinstance(result.get("scores"), dict) else {}
    rows = []
    for key, value in scores.items():
        if key == "fit_score":
            continue
        rows.append(
            {
                "key": key,
                "label": str(key).replace("_", " ").title(),
                "score": value,
                "weight": None,
                "explanation": "Legacy score component.",
            }
        )
    return rows


def _matched_skills(result: dict) -> list[dict]:
    matched = result.get("matched_skills")
    if isinstance(matched, list) and matched:
        return [_scrub_sensitive(item) for item in matched if isinstance(item, dict)]

    skills = result.get("skills", {}) if isinstance(result.get("skills"), dict) else {}
    out = []
    for requirement_type, key in (
        ("must_have", "matched_must_groups"),
        ("nice_to_have", "matched_nice_groups"),
    ):
        for item in skills.get(key, []) or []:
            if not isinstance(item, dict):
                continue
            group = item.get("group") or []
            skill = item.get("matched_by") or _group_label(group)
            out.append(
                {
                    "skill": skill,
                    "requirement_type": requirement_type,
                    "confidence": None,
                    "jd_requirement": _group_label(group),
                    "cv_evidence_ids": [],
                    "jd_evidence_ids": [],
                }
            )
    return out


def _missing_skills(result: dict) -> list[dict]:
    missing = result.get("missing_skills")
    if isinstance(missing, list) and missing:
        return [_scrub_sensitive(item) for item in missing if isinstance(item, dict)]

    gap = result.get("skill_gap", {}) if isinstance(result.get("skill_gap"), dict) else {}
    out = []
    for requirement_type, key, severity in (
        ("must_have", "missing_must_have", "high"),
        ("nice_to_have", "missing_nice_to_have", "medium"),
    ):
        for skill in gap.get(key, []) or []:
            out.append(
                {
                    "skill": skill,
                    "requirement_type": requirement_type,
                    "severity": severity,
                    "jd_requirement": skill,
                    "reason": f"JD mentions {skill}, but evidence was not found in the parsed CV.",
                    "suggestion": (
                        f"If you have actually used {skill}, add truthful project context and outcome. "
                        "Only add this if it is true."
                    ),
                }
            )
    return out


def _improvement_actions(result: dict) -> list[dict]:
    actions = result.get("improvement_actions")
    if isinstance(actions, list) and actions:
        return [_scrub_sensitive(item) for item in actions if isinstance(item, dict)]

    out = []
    gap = result.get("skill_gap", {}) if isinstance(result.get("skill_gap"), dict) else {}
    for index, item in enumerate(gap.get("learn_suggestions", []) or [], start=1):
        if not isinstance(item, dict):
            continue
        skill = item.get("skill") or "this skill"
        out.append(
            {
                "id": f"legacy_skill_{index:03d}",
                "priority": "medium",
                "title": f"Address {skill} evidence gap",
                "suggestion": (
                    f"If you have actually used {skill}, add a truthful CV bullet with project context. "
                    "Only add this if it is true."
                ),
                "guardrail": "Only add this if it is true. Do not invent skills or experience.",
            }
        )

    for index, item in enumerate(result.get("cv_improvements", []) or [], start=1):
        if not isinstance(item, dict):
            continue
        out.append(
            {
                "id": f"legacy_cv_{index:03d}",
                "priority": "medium",
                "title": item.get("issue") or "Improve CV evidence",
                "suggestion": (
                    f"If this reflects your actual work, {item.get('fix') or 'add clearer CV evidence'}. "
                    "Only add this if it is true."
                ),
                "guardrail": "Do not invent skills or experience.",
            }
        )
    return out


def _safe_rewrite_suggestions(result: dict) -> list[dict]:
    return _safe_dict_list(result.get("safe_rewrite_suggestions"))


def _interview_prep(result: dict) -> list[dict]:
    return _safe_dict_list(result.get("interview_prep"))


def _learning_roadmap(result: dict) -> list[dict]:
    return _safe_dict_list(result.get("learning_roadmap"))


def _has_v3_report_sections(result: dict) -> bool:
    if _safe_text(result.get("schema_version")).strip() == "3.0":
        return True

    for key in ("safe_rewrite_suggestions", "interview_prep", "learning_roadmap"):
        if key in result:
            return True

    actions = result.get("improvement_actions")
    if isinstance(actions, list):
        for item in actions:
            if not isinstance(item, dict):
                continue
            if any(
                key in item
                for key in (
                    "safe_suggestion",
                    "linked_skill",
                    "linked_evidence",
                    "status",
                    "do_not_fabricate",
                )
            ):
                return True

    limitations = result.get("limitations")
    if isinstance(limitations, list):
        limitation_text = " ".join(_safe_text(item).lower() for item in limitations)
        if "do not fabricate" in limitation_text or "not found in the parsed cv" in limitation_text:
            return True

    return False


def _limitations(result: dict) -> list[str]:
    limitations = result.get("limitations")
    out = [_safe_text(item) for item in limitations if item] if isinstance(limitations, list) else []
    required = [
        "This analysis estimates CV-to-JD fit only and does not guarantee any hiring outcome.",
        "Do not fabricate skills, experience, projects, employers, dates, certifications, or metrics.",
        "Missing evidence means support was not found in the parsed CV, not that the candidate definitely lacks the skill.",
    ]
    for item in required:
        if not any(item.lower() == existing.lower() for existing in out):
            out.append(item)
    if not out:
        out.append("The score is based on parsed CV/JD text and scoring heuristics.")
    return out


def _evidence(result: dict) -> list[dict]:
    evidence = result.get("evidence", [])
    if not isinstance(evidence, list):
        return []
    return [_scrub_sensitive(item) for item in evidence if isinstance(item, dict)]


def _evidence_by_id(result: dict) -> dict[str, dict]:
    return {
        str(item["id"]): item
        for item in _evidence(result)
        if item.get("id") is not None
    }


def _selected_evidence(evidence: list[dict]) -> list[dict]:
    cv_skill = [item for item in evidence if item.get("source") == "cv" and item.get("kind") == "skill"]
    cv_resp = [item for item in evidence if item.get("source") == "cv" and item.get("kind") == "responsibility"]
    jd_req = [item for item in evidence if item.get("source") == "jd" and item.get("kind") in {"requirement", "responsibility"}]
    return cv_skill[:4] + cv_resp[:4] + jd_req[:4]


def _evidence_snippet(item: dict) -> str:
    text = item.get("best_cv_bullet") or item.get("text") or item.get("jd_requirement") or "-"
    return _truncate(_safe_text(text), MAX_SNIPPET_CHARS)


def _extract_fit_score(result: dict) -> Any:
    scores = result.get("scores", {}) if isinstance(result.get("scores"), dict) else {}
    if scores.get("fit_score") is not None:
        return scores.get("fit_score")
    if result.get("fit_score") is not None:
        return result.get("fit_score")
    overall = result.get("overall", {}) if isinstance(result.get("overall"), dict) else {}
    return overall.get("fit_score")


def _fit_level(score: Any) -> str:
    try:
        value = float(score)
    except (TypeError, ValueError):
        value = 0.0
    if value >= 85:
        return "excellent"
    if value >= 70:
        return "good"
    if value >= 50:
        return "partial"
    return "weak"


def _format_score(value: Any) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.1f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return _safe_text(value)


def _format_weight(value: Any) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value) * 100:.0f}%"
    except (TypeError, ValueError):
        return _safe_text(value)


def _group_label(group: Any) -> str:
    if isinstance(group, list):
        return " / ".join(_safe_text(item) for item in group)
    return _safe_text(group)


def _scrub_sensitive(value: Any) -> Any:
    if isinstance(value, dict):
        return {
            key: _scrub_sensitive(item)
            for key, item in value.items()
            if str(key).lower() not in SENSITIVE_KEYS
        }
    if isinstance(value, list):
        return [_scrub_sensitive(item) for item in value]
    if isinstance(value, str):
        return _safe_text(value)
    return value


def _safe_dict_list(value: Any) -> list[dict]:
    if not isinstance(value, list):
        return []
    return [_scrub_sensitive(item) for item in value if isinstance(item, dict)]


def _safe_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [_safe_text(item) for item in value if item is not None]


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    for key in SENSITIVE_KEYS:
        text = re.sub(re.escape(key), "[redacted]", text, flags=re.IGNORECASE)
    for pattern in SENSITIVE_PATTERNS:
        text = pattern.sub("[redacted]", text)
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _truncate(value: str, max_chars: int) -> str:
    if len(value) <= max_chars:
        return value
    return value[: max_chars - 3].rstrip() + "..."
