from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace
import uuid

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.api.routes import jobs as jobs_route
from app.db.models import AnalysisJob
from app.db.session import get_db
from app.services.parsing.cv_parser import parse_cv_to_text
from app.services.parsing.jd_parser import parse_jd
from app.services.reporting.report_docx import build_docx_report
from app.services.scoring.result_v2 import build_result_v2
from app.services.scoring.result_v3 import build_result_v3
from app.services.scoring.scorer import score


def _docx_text(path_or_bytes) -> str:
    doc = Document(path_or_bytes)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _write_synthetic_cv(path):
    doc = Document()
    doc.add_heading("Synthetic Backend Candidate", level=1)
    doc.add_heading("Summary", level=2)
    doc.add_paragraph("Backend engineer using Python, FastAPI, Redis, Docker, and SQL.")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph(
        "Built FastAPI services and background workers for internal users with Docker deployment."
    )
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, Redis, Docker, SQL")
    doc.save(path)


def _patched_embeddings(texts):
    vectors = []
    for text in texts:
        lower = text.lower()
        if "api" in lower or "fastapi" in lower:
            vectors.append([1.0, 0.0, 0.0])
        elif "worker" in lower or "background" in lower:
            vectors.append([0.0, 1.0, 0.0])
        else:
            vectors.append([0.0, 0.0, 1.0])
    return vectors


def _build_worker_compatible_result(monkeypatch, tmp_path):
    import app.services.scoring.scorer as scorer_module

    cv_path = tmp_path / "candidate.docx"
    _write_synthetic_cv(cv_path)
    monkeypatch.setattr(scorer_module, "embed_texts", _patched_embeddings)

    cv_parsed = parse_cv_to_text(str(cv_path))
    jd_struct = parse_jd(
        "\n".join(
            [
                "Backend Engineer",
                "Required: Python and FastAPI experience.",
                "Required: AWS cloud operations.",
                "Nice to have: Docker or Kubernetes deployment experience.",
                "Build APIs and troubleshoot background workers.",
            ]
        )
    )
    jd_struct["must_have_skill_groups"].append(["aws"])
    jd_struct["skill_group_details"].append(
        {
            "group": ["aws"],
            "type": "required",
            "source_line": "Required: AWS cloud operations.",
        }
    )
    scored = score(cv_parsed, jd_struct)
    legacy_result = {
        "job_id": "job-1",
        "cv": {
            "file_name": "candidate.docx",
            "parsed_confidence": cv_parsed["confidence"],
            "skills_detected": cv_parsed.get("skills_detected", []),
        },
        "jd": jd_struct,
        **scored,
    }
    return build_result_v2(
        legacy_result,
        cv_parsed=cv_parsed,
        jd_struct=jd_struct,
        job_id="job-1",
    )


def test_worker_compatible_result_v2_and_docx_report(monkeypatch, tmp_path):
    result = _build_worker_compatible_result(monkeypatch, tmp_path)
    report_path = tmp_path / "report.docx"

    build_docx_report(result, str(report_path))

    assert result["schema_version"] == "2.0"
    for key in (
        "fit_score",
        "score_breakdown",
        "matched_skills",
        "missing_skills",
        "evidence",
        "improvement_actions",
        "limitations",
    ):
        assert result[key]
    assert result["scores"]["fit_score"] == result["fit_score"]
    assert result["overall"]["fit_score"] == result["fit_score"]
    assert report_path.exists()
    assert report_path.stat().st_size > 1000

    text = _docx_text(report_path)
    for expected in (
        "AI CV Fit Report",
        "Executive Summary",
        "Score Breakdown",
        "Matched Skills",
        "Missing Skills / Gaps",
        "Improvement Actions",
        "Limitations",
    ):
        assert expected in text
    for unsafe in (
        "access_token",
        "access_token_hash",
        "storage_path",
        "report_docx_path",
        "s3_key",
        "local_path",
        "raw_cv_text",
    ):
        assert unsafe not in text


def test_v2_result_api_and_report_download_compatibility(monkeypatch, tmp_path):
    result = _build_worker_compatible_result(monkeypatch, tmp_path)
    report_path = tmp_path / "report-api.docx"
    build_docx_report(result, str(report_path))
    report_bytes = report_path.read_bytes()

    job_id = uuid.uuid4()
    access_token = "correct-token"
    job = SimpleNamespace(
        id=job_id,
        status="succeeded",
        progress=100,
        error_message=None,
        result_json=result,
        report_docx_path="reports/report-api.docx",
        access_token_hash=jobs_route._hash_access_token(access_token),
    )
    fake_db = SimpleNamespace(get=lambda model, key: job if model is AnalysisJob else None)
    fake_storage = SimpleNamespace(read_bytes=lambda path: report_bytes)

    app = FastAPI()
    app.include_router(jobs_route.router)
    app.dependency_overrides[get_db] = lambda: fake_db
    monkeypatch.setattr(jobs_route, "get_storage", lambda: fake_storage)
    client = TestClient(app)

    result_response = client.get(f"/v1/jobs/{job_id}/result?access_token={access_token}")

    assert result_response.status_code == 200
    payload = result_response.json()
    nested = payload["result"]
    assert payload["overall_fit_score"] == result["fit_score"]
    assert payload["summary"]
    assert payload["evidence"]
    assert nested["schema_version"] == "2.0"
    assert nested["scores"]["fit_score"] == result["fit_score"]
    assert nested["fit_score"] == result["fit_score"]
    assert nested["overall"]["fit_score"] == result["fit_score"]
    assert nested["score_breakdown"]
    assert nested["matched_skills"]
    assert nested["missing_skills"]
    assert nested["improvement_actions"]
    assert nested["limitations"]

    download_response = client.get(f"/v1/jobs/{job_id}/report/download?access_token={access_token}")

    assert download_response.status_code == 200
    assert (
        download_response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(download_response.content) > 1000
    downloaded_text = _docx_text(BytesIO(download_response.content))
    assert "AI CV Fit Report" in downloaded_text
    assert "Executive Summary" in downloaded_text
    assert "Score Breakdown" in downloaded_text


def test_v3_result_report_download_includes_phase4_sections(monkeypatch, tmp_path):
    result = build_result_v3(_build_worker_compatible_result(monkeypatch, tmp_path))
    report_path = tmp_path / "report-v3-api.docx"
    build_docx_report(result, str(report_path))
    report_bytes = report_path.read_bytes()

    job_id = uuid.uuid4()
    access_token = "correct-token"
    job = SimpleNamespace(
        id=job_id,
        status="succeeded",
        progress=100,
        error_message=None,
        result_json=result,
        report_docx_path="reports/report-v3-api.docx",
        access_token_hash=jobs_route._hash_access_token(access_token),
    )
    fake_db = SimpleNamespace(get=lambda model, key: job if model is AnalysisJob else None)
    fake_storage = SimpleNamespace(read_bytes=lambda path: report_bytes)

    app = FastAPI()
    app.include_router(jobs_route.router)
    app.dependency_overrides[get_db] = lambda: fake_db
    monkeypatch.setattr(jobs_route, "get_storage", lambda: fake_storage)
    client = TestClient(app)

    download_response = client.get(f"/v1/jobs/{job_id}/report/download?access_token={access_token}")

    assert download_response.status_code == 200
    downloaded_text = _docx_text(BytesIO(download_response.content))
    assert "Improvement Action Plan" in downloaded_text
    assert "Safe Rewrite Suggestions" in downloaded_text
    assert "Interview Prep" in downloaded_text
    assert "Learning Roadmap" in downloaded_text
    assert "does not guarantee any hiring outcome" in downloaded_text
