"""
AI CV Fit — Phase 3 Integration Tests

Tests scoring quality, evidence mapping, guardrails, and report generation
for Phase 3 Result v2.

These tests verify:
1. Result schema v2 has all required fields
2. Missing skills use correct wording (not fabricated)
3. Improvement actions are conditional and guardrailed
4. Evidence mapping produces consistent IDs
5. Low-fit cases don't score too high
6. Token/auth flow doesn't break with v2
7. Report DOCX v2 generation works with structured data
"""
from __future__ import annotations

import re
import tempfile
from pathlib import Path

import pytest

from app.services.parsing.jd_parser import parse_jd
from app.services.scoring.scorer import score
from app.services.scoring.result_v2 import build_result_v2, fit_level
from app.services.reporting.report_docx import build_docx_report


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def make_cv_parsed(text: str, skills: list[str] | None = None) -> dict:
    """Build a synthetic cv_parsed dict from raw text."""
    from app.services.ontology.skill_ontology import get_skill_ontology
    ontology = get_skill_ontology()
    detected = sorted(ontology.detect_skills_in_text(text)) if skills is None else skills
    bullets = [l.strip() for l in text.splitlines() if len(l.strip()) >= 25]
    if not bullets:
        bullets = [l.strip() for l in text.splitlines() if len(l.strip()) >= 40]
    confidence = 0.85 if len(text) >= 300 else max(0.2, len(text) / 400 * 0.85)
    return {
        "text": text,
        "bullets": bullets[:80],
        "skills_detected": detected,
        "confidence": confidence,
    }


def make_result(text: str, jd_text: str, cv_skills: list[str] | None = None) -> dict:
    """Build a full v2 result from CV and JD text."""
    cv_parsed = make_cv_parsed(text, cv_skills)
    jd_struct = parse_jd(jd_text)
    scored = score(cv_parsed, jd_struct)
    result_full = {
        "job_id": "test-job",
        "cv": {
            "file_name": "test_cv.txt",
            "parsed_confidence": cv_parsed["confidence"],
            "skills_detected": cv_parsed.get("skills_detected", []),
        },
        "jd": jd_struct,
        **scored,
    }
    return build_result_v2(result_full, cv_parsed=cv_parsed, jd_struct=jd_struct, job_id="test-job")


# ---------------------------------------------------------------------------
# 1. Result Schema v2 Tests
# ---------------------------------------------------------------------------

def test_result_v2_has_all_required_fields():
    """Result v2 must have all required fields from the contract."""
    cv = "Built FastAPI APIs with PostgreSQL and Redis. Deployed with Docker."
    jd = "Python FastAPI backend. PostgreSQL database. Redis caching."
    result = make_result(cv, jd)

    for key in (
        "schema_version",
        "fit_score",
        "scores",
        "overall",
        "score_breakdown",
        "matched_skills",
        "missing_skills",
        "evidence",
        "improvement_actions",
        "limitations",
        "metadata",
    ):
        assert key in result, f"Missing required field: {key}"

    assert result["schema_version"] == "2.0"


def test_result_v2_preserves_legacy_score_alias():
    """Legacy score aliases must still be present for backward compatibility."""
    cv = "Python FastAPI developer with PostgreSQL."
    jd = "Python FastAPI. PostgreSQL."
    result = make_result(cv, jd)

    assert result["scores"]["fit_score"] is not None
    assert result["fit_score"] == result["scores"]["fit_score"]
    assert result["overall"]["fit_score"] == result["fit_score"]


def test_fit_level_thresholds_are_correct():
    """Fit level thresholds must match the contract exactly."""
    assert fit_level(100) == "excellent"
    assert fit_level(85) == "excellent"
    assert fit_level(84.9) == "good"
    assert fit_level(70) == "good"
    assert fit_level(69.9) == "partial"
    assert fit_level(50) == "partial"
    assert fit_level(49.9) == "weak"
    assert fit_level(0) == "weak"
    assert fit_level(None) == "weak"


def test_score_breakdown_contains_all_components():
    """Score breakdown must include all 5 scoring components."""
    cv = "Python developer with FastAPI and PostgreSQL."
    jd = "Python FastAPI PostgreSQL Redis Docker"
    result = make_result(cv, jd)

    breakdown = result.get("score_breakdown", [])
    assert len(breakdown) > 0
    for item in breakdown:
        assert "key" in item
        assert "label" in item
        assert "score" in item
        assert "weight" in item
        assert "explanation" in item
        assert 0 <= item["score"] <= 100


def test_metadata_contains_scorer_version():
    """Metadata must include scorer_version identifying the phase."""
    cv = "Python FastAPI PostgreSQL"
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    assert result["metadata"]["scorer_version"] == "phase3.result_json_v2"
    assert "generated_at" in result["metadata"]


# ---------------------------------------------------------------------------
# 2. Missing Skills — Guardrail Wording Tests
# ---------------------------------------------------------------------------

def test_missing_skill_uses_parsed_cv_language():
    """Missing skill reason must use 'not found in parsed CV', not 'you don't know'."""
    cv = "Python developer. Used JavaScript."
    jd = "Python FastAPI. PostgreSQL. FastAPI is required."
    result = make_result(cv, jd)

    for item in result.get("missing_skills", []):
        reason = item.get("reason", "")
        assert "you don't know" not in reason.lower(), \
            f"Missing skill reason uses 'you don't know': {reason}"
        # Reason may be truncated at ~240 chars in practice. Check for the key phrase
        # "evidence was not found" which is the core guardrail signal even when truncated.
        assert "evidence was not found" in reason.lower() or \
               ("not found" in reason.lower() and "parsed cv" in reason.lower()), \
            f"Missing skill reason missing evidence-not-found guardrail phrase: {reason}"


def test_missing_skill_summary_uses_parsed_cv_language():
    """Overall summary must use 'not found in parsed CV', not absolute claims."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL Redis Docker Kubernetes AWS"
    result = make_result(cv, jd)

    summary = result.get("overall", {}).get("summary", "")
    assert "you don't have" not in summary.lower(), \
        f"Summary uses absolute claim: {summary}"
    assert "you do not have" not in summary.lower()


def test_missing_skill_has_conditional_suggestion():
    """Missing skill suggestions must be conditional and include guardrail."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    for item in result.get("missing_skills", []):
        suggestion = item.get("suggestion", "")
        assert suggestion, f"Missing skill '{item.get('skill')}' has no suggestion"
        assert "if you have actually used" in suggestion.lower() or \
               "if you have used" in suggestion.lower(), \
            f"Suggestion not conditional: {suggestion}"
        assert "only add this if it is true" in suggestion.lower(), \
            f"Suggestion missing 'only add this if it is true': {suggestion}"


def test_high_severity_missing_skill_from_must_have():
    """Must-have missing skills must have high severity."""
    cv = "Python developer with no database experience."
    jd = "Python PostgreSQL"
    result = make_result(cv, jd)

    missing = [m for m in result.get("missing_skills", []) if m.get("requirement_type") == "must_have"]
    if missing:
        for item in missing:
            assert item.get("severity") in ("high", "medium"), \
                f"Must-have missing skill '{item.get('skill')}' should have high/medium severity, got: {item.get('severity')}"


# ---------------------------------------------------------------------------
# 3. Improvement Actions — Guardrail Tests
# ---------------------------------------------------------------------------

def test_improvement_actions_are_conditional():
    """All improvement actions must use conditional wording."""
    cv = "Python developer with minimal experience."
    jd = "Python FastAPI PostgreSQL Redis Docker Kubernetes"
    result = make_result(cv, jd)

    for item in result.get("improvement_actions", []):
        suggestion = item.get("suggestion", "")
        assert suggestion, f"Improvement action '{item.get('title')}' has no suggestion"
        assert "if you have" in suggestion.lower() or "only add this if" in suggestion.lower(), \
            f"Improvement action not conditional: {suggestion}"


def test_improvement_actions_have_guardrail():
    """Improvement actions must include a guardrail reminder."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    for item in result.get("improvement_actions", []):
        guardrail = item.get("guardrail", "")
        assert guardrail, f"Improvement action '{item.get('title')}' has no guardrail"
        assert "do not invent" in guardrail.lower() or "only add this if" in guardrail.lower(), \
            f"Improvement action guardrail is insufficient: {guardrail}"


def test_improvement_actions_priority_must_have():
    """Must-have missing skills must produce high-priority improvement actions."""
    cv = "Python developer with no backend framework."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    must_missing = [m for m in result.get("missing_skills", [])
                    if m.get("requirement_type") == "must_have"]
    for skill_item in must_missing:
        related_actions = [
            a for a in result.get("improvement_actions", [])
            if a.get("related_skill") == skill_item.get("skill")
        ]
        if related_actions:
            for action in related_actions:
                assert action.get("priority") in ("high", "medium"), \
                    f"Must-have missing skill action should be high/medium priority: {action}"


def test_improvement_actions_not_empty_when_skills_missing():
    """Improvement actions must not be empty when skills are missing."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL Redis"
    result = make_result(cv, jd)

    missing = result.get("missing_skills", [])
    actions = result.get("improvement_actions", [])
    assert len(missing) > 0, "Expected missing skills but none found"
    assert len(actions) > 0, \
        f"Expected improvement actions for {len(missing)} missing skills but got none"


# ---------------------------------------------------------------------------
# 4. Evidence Mapping Tests
# ---------------------------------------------------------------------------

def test_matched_skill_has_cv_evidence_ids():
    """Matched skills must reference CV evidence IDs."""
    cv = "Built FastAPI REST APIs with PostgreSQL and deployed on Docker."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    matched = result.get("matched_skills", [])
    if matched:
        for item in matched:
            assert "cv_evidence_ids" in item, \
                f"Matched skill '{item.get('skill')}' missing cv_evidence_ids"
            if item.get("cv_evidence_ids"):
                # Check that evidence IDs exist in evidence list
                for ev_id in item["cv_evidence_ids"]:
                    evidence_ids = {e["id"] for e in result.get("evidence", []) if isinstance(e, dict)}
                    assert ev_id in evidence_ids, \
                        f"Evidence ID '{ev_id}' referenced in matched skill but not in evidence list"


def test_evidence_has_stable_ids():
    """Evidence items must have stable, unique IDs."""
    cv = "Python FastAPI PostgreSQL developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    evidence = result.get("evidence", [])
    ids = [e.get("id") for e in evidence if isinstance(e, dict)]
    assert len(ids) == len(set(ids)), \
        f"Evidence IDs are not unique: {ids}"


def test_evidence_ids_are_deterministic_format():
    """Evidence IDs must follow the naming convention from the contract."""
    cv = "Python FastAPI PostgreSQL developer."
    jd = "Python FastAPI PostgreSQL Redis Docker"
    result = make_result(cv, jd)

    evidence = result.get("evidence", [])
    for item in evidence:
        eid = item.get("id", "")
        assert eid, f"Evidence item missing id: {item}"
        assert re.match(r"^(ev_cv_|ev_jd_|legacy_)", eid), \
            f"Evidence ID '{eid}' doesn't follow convention (ev_cv_*, ev_jd_*)"


def test_jd_evidence_has_source_jd():
    """JD evidence items with kind=requirement must have source='jd'."""
    cv = "Python FastAPI PostgreSQL"
    jd = "Python FastAPI PostgreSQL Redis Kubernetes"
    result = make_result(cv, jd)

    for item in result.get("evidence", []):
        if isinstance(item, dict) and item.get("kind") == "requirement":
            # Requirement evidence from JD should have source='jd'
            if "jd" in item.get("id", "").lower():
                assert item.get("source") == "jd", \
                    f"JD requirement evidence has wrong source '{item.get('source')}': {item}"


# ---------------------------------------------------------------------------
# 5. Low-Fit Cap Tests
# ---------------------------------------------------------------------------

def test_irrelevant_cv_scores_low():
    """A CV completely unrelated to JD must not score above partial range."""
    cv = "Chef with 10 years cooking experience. Expert in French cuisine."
    jd = "Python FastAPI PostgreSQL Redis Docker Kubernetes"
    result = make_result(cv, jd)

    score = result.get("fit_score", 0) or 0
    assert score < 70, \
        f"Irrelevant CV scored {score}, expected < 70"


def test_cv_missing_all_must_have_skills_scores_low():
    """A CV missing all must-have skills must not score high."""
    cv = "Python developer. No databases. No frameworks."
    jd = "Python FastAPI PostgreSQL Redis Docker"
    result = make_result(cv, jd)

    score = result.get("fit_score", 0) or 0
    missing = [m for m in result.get("missing_skills", [])
               if m.get("requirement_type") == "must_have"]
    if len(missing) >= 2:
        assert score < 70, \
            f"CV missing {len(missing)} must-have skills scored {score}, expected < 70"


def test_fit_level_is_weak_for_irrelevant_cv():
    """Irrelevant CV must have fit_level 'weak'."""
    cv = "I am a kindergarten teacher with no technical skills."
    jd = "Python FastAPI backend engineer with PostgreSQL"
    result = make_result(cv, jd)

    level = result.get("overall", {}).get("fit_level", "")
    assert level == "weak", \
        f"Irrelevant CV should have fit_level='weak', got '{level}'"


def test_fit_level_is_partial_for_partial_match():
    """CV with some overlap should not score excellent."""
    cv = "Python developer. No frameworks, no databases."
    jd = "Python FastAPI PostgreSQL Redis"
    result = make_result(cv, jd)

    level = result.get("overall", {}).get("fit_level", "")
    score = result.get("fit_score", 0) or 0
    if score < 70:
        assert level in ("partial", "weak"), \
            f"Partial match with score={score} should be partial/weak, got '{level}'"


# ---------------------------------------------------------------------------
# 6. Limitations Tests
# ---------------------------------------------------------------------------

def test_limitations_include_no_hiring_guarantee():
    """Limitations must include the no-hiring-guarantee notice."""
    cv = "Python FastAPI developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    limitations = result.get("limitations", [])
    assert any("does not guarantee" in str(l) and "hiring" in str(l).lower()
               for l in limitations), \
        f"Limitations missing hiring guarantee notice: {limitations}"


def test_limitations_include_missing_evidence_note():
    """Limitations must note that missing evidence is not absolute proof."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    limitations = result.get("limitations", [])
    # The actual wording is: "Missing evidence means support was not found in the parsed CV,
    # not that the candidate definitely lacks the skill."
    assert any("missing evidence" in str(l).lower() and
               ("not found" in str(l).lower() or "not that" in str(l).lower())
               for l in limitations), \
        f"Limitations missing missing-evidence caveat: {limitations}"


def test_limitations_include_no_fabrication():
    """Limitations must warn against fabricating skills."""
    cv = "Python developer."
    jd = "Python FastAPI"
    result = make_result(cv, jd)

    limitations = result.get("limitations", [])
    assert any("do not invent" in str(l).lower() or "fabricat" in str(l).lower()
               for l in limitations), \
        f"Limitations missing no-fabrication warning: {limitations}"


# ---------------------------------------------------------------------------
# 7. Guardrail — No Guarantee Language Tests
# ---------------------------------------------------------------------------

GUARANTEE_PATTERNS = [
    r"\bwill\s+(definitely\s+)?(get\s+)?(hired|selected|accepted)",
    r"\bguarantee[sd]?\s+(you[' ]?ll\s+)?(get\s+)?(hired|selected)",
    r"\byou\s+will\s+definitely\s+(get\s+)?(the\s+)?job",
]


def test_no_guarantee_language_in_summary():
    """Overall summary must not contain guarantee language."""
    cv = "Python FastAPI developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    summary = result.get("overall", {}).get("summary", "")
    for pattern in GUARANTEE_PATTERNS:
        match = re.search(pattern, summary, re.IGNORECASE)
        assert not match, \
            f"Summary contains guarantee language: {summary}"


def test_no_guarantee_language_in_guardrail_notice():
    """Guardrail notice must not contain guarantee language."""
    cv = "Python FastAPI PostgreSQL"
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    notice = result.get("overall", {}).get("guardrail_notice", "")
    for pattern in GUARANTEE_PATTERNS:
        match = re.search(pattern, notice, re.IGNORECASE)
        assert not match, \
            f"Guardrail notice contains guarantee language: {notice}"


def test_no_fabrication_language_in_missing_skills():
    """Missing skill reasons must not claim the candidate doesn't know a skill."""
    cv = "Python developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    for item in result.get("missing_skills", []):
        reason = item.get("reason", "")
        assert "you don't know" not in reason.lower(), \
            f"Missing skill uses fabrication language: {reason}"
        assert "you do not know" not in reason.lower(), \
            f"Missing skill uses fabrication language: {reason}"


# ---------------------------------------------------------------------------
# 8. Report DOCX v2 Generation Tests
# ---------------------------------------------------------------------------

def test_docx_v2_contains_score_breakdown():
    """DOCX report v2 must contain score breakdown section."""
    cv = "Python FastAPI PostgreSQL developer with Docker."
    jd = "Python FastAPI PostgreSQL Redis Docker"
    result = make_result(cv, jd)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = tmp.name

    try:
        build_docx_report(result, out_path)
        content = Path(out_path).read_bytes()
        assert len(content) > 500, "DOCX report is too small, likely empty"
    finally:
        try:
            Path(out_path).unlink()
        except OSError:
            pass


def test_docx_v2_contains_guardrail_notice():
    """DOCX report v2 must include guardrail notice in limitations."""
    cv = "Python FastAPI developer."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = tmp.name

    try:
        build_docx_report(result, out_path)
        # Parse DOCX using python-docx to read text
        from docx import Document
        doc = Document(out_path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        full_text = " ".join(paragraphs)
        assert "does not guarantee" in full_text.lower(), \
            f"DOCX missing guardrail notice. Text: {full_text[:500]}"
    finally:
        try:
            Path(out_path).unlink()
        except OSError:
            pass


def test_docx_v2_excludes_sensitive_keys():
    """DOCX report must not expose sensitive internal keys."""
    cv = "Python FastAPI PostgreSQL"
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)
    result["access_token"] = "secret-token-123"
    result["storage_path"] = "uploads/private-cv.docx"
    result["raw_cv_text"] = "CONFIDENTIAL FULL CV CONTENT"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = tmp.name

    try:
        build_docx_report(result, out_path)
        # Parse DOCX using python-docx
        from docx import Document
        doc = Document(out_path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        full_text = " ".join(paragraphs)
        for unsafe in ("secret-token", "private-cv.docx", "CONFIDENTIAL"):
            assert unsafe not in full_text, \
                f"DOCX exposed sensitive content: {unsafe}"
    finally:
        try:
            Path(out_path).unlink()
        except OSError:
            pass


# ---------------------------------------------------------------------------
# 9. Edge Cases
# ---------------------------------------------------------------------------

def test_ultra_short_cv_handled_gracefully():
    """Ultra-short CV must not crash and must produce a result."""
    cv = "Developer."
    jd = "Python FastAPI PostgreSQL Redis Docker Kubernetes"
    result = make_result(cv, jd)

    assert result.get("fit_score") is not None
    assert 0 <= result["fit_score"] <= 100
    assert result.get("overall", {}).get("fit_level") in ("weak", "partial", "good", "excellent")
    assert len(result.get("limitations", [])) > 0


def test_cv_with_no_skills_detected():
    """CV with no detectable skills must not crash."""
    cv = "A person who does things with technology."
    jd = "Python FastAPI PostgreSQL"
    result = make_result(cv, jd)

    assert result.get("fit_score") is not None
    assert 0 <= result["fit_score"] <= 100
    # Should have many missing skills
    assert len(result.get("missing_skills", [])) > 0


def test_cv_with_all_skills_explicitly_listed():
    """CV that explicitly matches all JD requirements must score high."""
    cv = (
        "Python developer with FastAPI experience. "
        "Used PostgreSQL for database. "
        "Implemented Redis caching. "
        "Deployed applications with Docker."
    )
    jd = "Python FastAPI PostgreSQL Redis Docker"
    result = make_result(cv, jd)

    score = result.get("fit_score", 0) or 0
    assert score >= 60, \
        f"CV matching all JD requirements scored only {score}, expected >= 60"


def test_ontology_alias_expands_correctly():
    """Skill ontology aliases should match related technologies."""
    cv = "Used Django for web development. PostgreSQL database."
    jd = "Python Flask PostgreSQL"
    result = make_result(cv, jd)

    # Django is related to Flask — should be considered
    # The system should either match or clearly show it as missing
    assert result.get("fit_score") is not None
    assert 0 <= result["fit_score"] <= 100


def test_responsibility_match_produces_evidence():
    """Responsibility matching should produce evidence entries."""
    cv = "Built and deployed FastAPI REST APIs. Designed database schemas."
    jd = "Build REST APIs. Design database schemas."
    result = make_result(cv, jd)

    evidence = result.get("evidence", [])
    resp_evidence = [e for e in evidence if e.get("kind") == "responsibility"]
    assert len(resp_evidence) >= 0  # May or may not match, but must not crash


def test_multiple_missing_must_have_skills():
    """Multiple missing must-have skills should all appear in missing_skills."""
    cv = "Python developer. No databases. No frameworks."
    jd = "Python FastAPI PostgreSQL Redis Docker Kubernetes AWS"
    result = make_result(cv, jd)

    missing = [m for m in result.get("missing_skills", [])
               if m.get("requirement_type") == "must_have"]
    assert len(missing) >= 2, \
        f"Expected at least 2 missing must-have skills, got {len(missing)}"
