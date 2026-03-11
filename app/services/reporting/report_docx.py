from docx import Document


def build_docx_report(result: dict, out_path: str):
    doc = Document()
    doc.add_heading("CV Fit Report (CV vs JD) - v2", level=1)

    cv_meta = result.get("cv", {})
    jd_meta = result.get("jd", {})
    scores = result.get("scores", {})

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(f"File: {cv_meta.get('file_name', '-')}")
    doc.add_paragraph(f"Parsed confidence: {cv_meta.get('parsed_confidence', '-')}")
    doc.add_paragraph(f"Target role: {jd_meta.get('role', '-') or '-'}")
    doc.add_paragraph(f"Fit score: {scores.get('fit_score', '-')}")

    doc.add_heading("Score Breakdown", level=2)
    for k, v in scores.items():
        doc.add_paragraph(f"{k}: {v}")

    gap = result.get("skill_gap", {})
    doc.add_heading("Skill Gap Analysis", level=2)
    doc.add_paragraph("Missing must-have: " + (", ".join(gap.get("missing_must_have", [])) or "None"))
    doc.add_paragraph("Missing nice-to-have: " + (", ".join(gap.get("missing_nice_to_have", [])) or "None"))

    doc.add_heading("Learning Plan", level=2)
    for item in gap.get("learn_suggestions", []):
        doc.add_paragraph(
            f"- {item['skill']}: {item['reason']} | "
            f"level={item['resources_level']} | "
            f"estimated={item.get('time_estimate_weeks', '-') } weeks"
        )

    doc.add_heading("CV Improvements", level=2)
    for item in result.get("cv_improvements", []):
        doc.add_paragraph(f"- {item['issue']} → {item['fix']}")

    doc.add_heading("Evidence", level=2)
    for e in result.get("evidence", []):
        if e.get("type") == "responsibility_match":
            doc.add_paragraph(
                f"- JD: {e.get('jd_requirement','')}\n"
                f"  CV: {e.get('text','')}\n"
                f"  similarity={e.get('similarity','')}"
            )
        else:
            doc.add_paragraph(f"- {e.get('text','')}")

    doc.save(out_path)