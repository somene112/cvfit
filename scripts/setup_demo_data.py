"""
Phase 5 Demo Data Setup Script (v2 — fixed API contracts)
"""
from __future__ import annotations
import json, sys, uuid, time, os, io
from pathlib import Path
from typing import Optional

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "backend"))

import urllib.request, urllib.error

# ── Config ──────────────────────────────────────────────────────────────────
BASE_URL = os.environ.get("API_BASE_URL", "https://cvfit.onrender.com")
DEMO_EMAIL = "dat_phase5_demo@demo.app"
DEMO_PASS  = "DemoTest123!"
DEMO_NAME  = "Dat Demo"

SAMPLE_JD = (
    "Backend Engineer\n"
    "Company: DemoCorp\n\n"
    "Requirements:\n"
    "- 2+ years Python, FastAPI or Django\n"
    "- PostgreSQL and Redis experience\n"
    "- Docker, Kubernetes\n"
    "- CI/CD with GitHub Actions\n"
    "- REST API design\n"
    "- TypeScript is a plus\n"
)


def _make_cv_docx() -> bytes:
    """Build a minimal DOCX containing the synthetic CV text."""
    try:
        from docx import Document
    except ImportError:
        return b"PK\x03\x04"  # placeholder — won't pass upload validation anyway

    doc = Document()
    doc.add_heading("Nguyen Van A — Software Engineer", 0)
    doc.add_paragraph("Email: nguyen.van.a@example.com  |  Phone: +84-123-456-7890")
    doc.add_paragraph("")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Backend Engineer with 2+ years building high-traffic REST APIs using FastAPI, "
        "PostgreSQL, and Redis. Improved API latency by 40% at TechCorp."
    )
    doc.add_heading("Skills", level=1)
    for skill in ["Python", "FastAPI", "PostgreSQL", "Redis", "Docker",
                  "Kubernetes", "GitHub Actions", "TypeScript", "AWS EKS"]:
        doc.add_paragraph(skill, style="List Bullet")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(
        "Software Engineer — TechCorp (2022–Present)\n"
        "Built REST APIs with FastAPI; managed PostgreSQL + Redis; deployed with "
        "Docker and Kubernetes on AWS EKS; implemented CI/CD with GitHub Actions. "
        "Improved API latency by 40% via Redis caching and query optimisation."
    )
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(
        "CV Fit App: AI-powered CV-job matching platform. FastAPI backend, "
        "Next.js frontend, Celery workers, PostgreSQL + Redis, deployed on Render.\n\n"
        "E-Commerce Product API (FastAPI/PostgreSQL/Redis): 10k req/day, 99.9% uptime."
    )
    doc.add_heading("Education", level=1)
    doc.add_paragraph("BSc Computer Science, University of Tech, 2022")
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Cloud Practitioner (2023)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ─────────────────────────────────────────────────────────────────
class C:
    OK, ERR, WRN, INF, END = "\033[92m", "\033[91m", "\033[93m", "\033[94m", "\033[0m"

def log(status: str, msg: str):
    color = {"PASS": C.OK, "FAIL": C.ERR, "WARN": C.WRN, "INFO": C.INF}.get(status, C.END)
    print(f"  [{color}{status}{C.END}] {msg}")

def api_call(method: str, path: str, token: str = None,
             payload: dict = None, file_data: bytes = None,
             file_name: str = None, file_field: str = None,
             timeout: int = 30) -> tuple[int, dict | list | str, str]:
    url = BASE_URL + path
    try:
        if file_data is not None:
            boundary = uuid.uuid4().hex
            parts = []
            parts.append(f"--{boundary}\r\nContent-Disposition: form-data; "
                          f'name="{file_field}"; filename="{file_name}"\r\n'
                          f"Content-Type: application/octet-stream\r\n\r\n".encode())
            parts.append(file_data)
            parts.append(b"\r\n")
            parts.append(f"--{boundary}--\r\n".encode())
            body_bytes = b"".join(parts)
            req = urllib.request.Request(url, data=body_bytes, method=method)
            req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
        else:
            data = json.dumps(payload).encode() if payload else None
            req = urllib.request.Request(url, data=data, method=method)
            req.add_header("Content-Type", "application/json")
        if token:
            req.add_header("Authorization", f"Bearer {token}")
        with urllib.request.urlopen(req, timeout=timeout) as r:
            raw = r.read().decode(errors="replace")
            try:
                parsed = json.loads(raw)
            except Exception:
                parsed = raw
            return r.status, parsed, raw
    except urllib.error.HTTPError as e:
        raw = e.read().decode(errors="replace")
        try:
            parsed = json.loads(raw)
        except Exception:
            parsed = raw
        return e.code, parsed, raw
    except Exception as ex:
        return 0, {}, str(ex)


def wait_job(token: str, job_id: str, timeout: int = 150) -> str:
    log("INFO", f"Polling job {job_id[:8]}... (timeout {timeout}s)")
    deadline = time.time() + timeout
    while time.time() < deadline:
        code, body, _ = api_call("GET", f"/v1/jobs/{job_id}", token=token)
        if code == 200 and isinstance(body, dict):
            status = body.get("status", "")
            pct = body.get("progress", 0)
            log("INFO", f"  status={status}  progress={pct}%")
            if status in ("completed", "failed"):
                return status
        time.sleep(5)
    return "timeout"


# ── Step 0: Register + Login ─────────────────────────────────────────────────
def step0():
    log("INFO", "Step 0 — Register/login")
    code, body, _ = api_call("POST", "/v1/auth/register",
                              payload={"email": DEMO_EMAIL, "password": DEMO_PASS,
                                       "name": DEMO_NAME})
    if code in (200, 201):
        log("PASS", f"Registered: {DEMO_EMAIL}")
    elif code == 409:
        log("INFO", f"User exists — continuing")
    else:
        log("FAIL", f"Register {code}: {body}")

    code, body, _ = api_call("POST", "/v1/auth/login",
                              payload={"email": DEMO_EMAIL, "password": DEMO_PASS})
    if code != 200:
        log("FAIL", f"Login failed {code}: {body}")
        sys.exit(1)
    token = body.get("access_token", "") if isinstance(body, dict) else ""
    log("PASS", f"Logged in  token={token[:12]}...")
    return token


# ── Step 1: CV Upload (DOCX) ─────────────────────────────────────────────────
def step1(token: str) -> str | None:
    log("INFO", "Step 1 — Upload CV (DOCX)")
    cv_bytes = _make_cv_docx()
    code, body, _ = api_call("POST", "/v1/cv/upload", token=token,
                              file_data=cv_bytes,
                              file_name="cv_demo.docx",
                              file_field="file")
    if code not in (200, 201):
        log("FAIL", f"CV upload {code}: {body}")
        return None
    cv_id = (body.get("cv_id") or body.get("cv_file_id") or
             body.get("id") or "") if isinstance(body, dict) else ""
    log("PASS", f"CV uploaded  cv_id={cv_id[:8] if cv_id else '?'}...")
    return str(cv_id)


# ── Step 2: Create analysis job ──────────────────────────────────────────────
def step2(token: str, cv_id: str) -> tuple[str, str]:
    log("INFO", "Step 2 — Create analysis job")
    payload = {
        "cv_file_id": cv_id,
        "jd_text": SAMPLE_JD,
        "options": {"target_role": "Backend Engineer"},
    }
    code, body, _ = api_call("POST", "/v1/jobs/create-score", token=token, payload=payload)
    if code not in (200, 201):
        log("FAIL", f"Create score {code}: {body}")
        sys.exit(1)
    job_id = body.get("job_id", "") if isinstance(body, dict) else ""
    log("PASS", f"Job queued  job_id={job_id[:8] if job_id else '?'}")
    status = wait_job(token, job_id)
    log("PASS" if status == "completed" else "WARN",
        f"Job finished with status={status}")
    return job_id, status


# ── Step 3: Create application ───────────────────────────────────────────────
def step3(token: str) -> str:
    log("INFO", "Step 3 — Create application")
    payload = {"job_title": "Backend Engineer",
               "company_name": "DemoCorp", "jd_text": SAMPLE_JD}
    code, body, _ = api_call("POST", "/v1/applications", token=token, payload=payload)
    if code not in (200, 201):
        log("FAIL", f"Create application {code}: {body}")
        sys.exit(1)
    app_id = (body.get("id") or "") if isinstance(body, dict) else ""
    log("PASS", f"Application created  app_id={app_id[:8] if app_id else '?'}")
    return str(app_id)


# ── Step 4: Attach analysis ─────────────────────────────────────────────────
def step4(token: str, app_id: str, job_id: str):
    log("INFO", "Step 4 — Attach analysis")
    code, body, _ = api_call("POST",
                               f"/v1/applications/{app_id}/attach-analysis/{job_id}",
                               token=token)
    if code not in (200, 201):
        log("FAIL", f"Attach {code}: {body}")
        sys.exit(1)
    log("PASS", f"Analysis attached")


# ── Step 5: Cover letter ────────────────────────────────────────────────────
def step5(token: str, app_id: str):
    log("INFO", "Step 5 — Generate cover letter")
    code, body, _ = api_call("POST",
                               f"/v1/applications/{app_id}/cover-letter/generate",
                               token=token)
    if code not in (200, 201):
        log("FAIL", f"Cover letter {code}: {body}")
        return
    log("PASS", f"Cover letter generated")


# ── Step 6: Application package ─────────────────────────────────────────────
def step6(token: str, app_id: str):
    log("INFO", "Step 6 — Generate application package")
    code, body, _ = api_call("POST",
                               f"/v1/applications/{app_id}/package/generate",
                               token=token)
    if code not in (200, 201):
        log("FAIL", f"Package {code}: {body}")
        return
    if isinstance(body, dict):
        sections = list(body.keys())
        log("PASS", f"Package generated ({len(sections)} sections: {sections})")
    else:
        log("PASS", "Package generated")


# ── Step 7: Career profile ─────────────────────────────────────────────────
def step7(token: str) -> list:
    log("INFO", "Step 7 — Add career profile items")
    items = []
    for item_type, title, desc, evidence in [
        ("skill",      "Docker",       "Containerized microservices on AWS EKS",          "Used Docker Compose for local dev; EKS for production"),
        ("project",    "CV Fit App",   "AI-powered CV-job matching platform",              "Built end-to-end: FastAPI + Next.js + Celery + PostgreSQL"),
        ("achievement","40% Latency Reduction",
                                                "Improved API latency at TechCorp",         "40% latency reduction via Redis caching + PG query optimisation"),
    ]:
        payload = {
            "item_type": item_type,
            "title": title,
            "description": desc,
            "evidence_text": evidence,
        }
        code, body, _ = api_call("POST", "/v1/profile/items", token=token, payload=payload)
        status = "PASS" if code in (200, 201) else "FAIL"
        log(status, f"  {item_type}: {title}")
        if code in (200, 201):
            items.append(body)
    log("PASS", f"{len(items)} profile items added")
    return items


# ── Step 8: Interview practice ──────────────────────────────────────────────
def step8(token: str, app_id: str):
    log("INFO", "Step 8 — Interview practice")
    # Get questions
    code, body, _ = api_call("GET",
                               f"/v1/applications/{app_id}/interview/questions",
                               token=token)
    if code not in (200, 201):
        log("FAIL", f"Get questions {code}: {body}")
        return
    questions = (body.get("questions", []) if isinstance(body, dict)
                 else body if isinstance(body, list) else [])
    log("INFO", f"  Got {len(questions)} questions")
    if not questions:
        log("WARN", "  No questions returned — may need analysis attached")
        return

    for q in questions[:2]:
        q_id  = q.get("question_id") if isinstance(q, dict) else None
        q_txt = q.get("question")   if isinstance(q, dict) else str(q)
        if q_id:
            payload = {
                "question_id": q_id,
                "question":    q_txt,
                "answer_text": (
                    "I improved API latency by 40% at TechCorp using Redis caching "
                    "and PostgreSQL query optimisation. The system now handles 10k "
                    "requests per day with 99.9% uptime."
                ),
            }
            code2, body2, _ = api_call(
                "POST",
                f"/v1/applications/{app_id}/interview/answers",
                token=token, payload=payload,
            )
            s = "PASS" if code2 in (200, 201) else "FAIL"
            log(s, f"  Answer submitted for question {q_id[:8]}")

    # List answers
    code, body, _ = api_call("GET",
                               f"/v1/applications/{app_id}/interview/answers",
                               token=token)
    if code == 200:
        answers = (body.get("answers", []) if isinstance(body, dict)
                   else body if isinstance(body, list) else [])
        log("PASS", f"  {len(answers)} answer(s) in history")
    else:
        log("WARN", f"  List answers returned {code}")


# ── Step 9: Readiness + usage ───────────────────────────────────────────────
def step9(token: str, app_id: str):
    log("INFO", "Step 9 — Readiness + usage")
    code, body, _ = api_call("GET",
                               f"/v1/applications/{app_id}/readiness",
                               token=token)
    if code == 200:
        level = (body.get("readiness_level", "unknown")
                 if isinstance(body, dict) else "unknown")
        log("PASS", f"Readiness: {level}")
    else:
        log("WARN", f"Readiness {code}")

    code, body, _ = api_call("GET", "/v1/usage/me", token=token)
    if code == 200:
        jobs = body.get("jobs_created", 0) if isinstance(body, dict) else 0
        apps = body.get("applications_created", 0) if isinstance(body, dict) else 0
        log("PASS", f"Usage: jobs={jobs}, apps={apps}")
    else:
        log("WARN", f"Usage {code}")


# ── Cleanup ───────────────────────────────────────────────────────────────────
def cleanup(token: str):
    log("INFO", "Cleanup — deleting demo data")
    for path, key in [("/v1/applications", "items"),
                      ("/v1/jobs/history", "items"),
                      ("/v1/profile/items", "items")]:
        code, body, _ = api_call("GET", path, token=token)
        if code == 200:
            items = (body.get(key, []) if isinstance(body, dict)
                     else body if isinstance(body, list) else [])
            for item in items:
                rid = item.get("id") or item.get("job_id") or ""
                if rid:
                    p = path.replace("/history", "").replace("/items", f"/items/{rid}").replace("/v1/applications", f"/v1/applications/{rid}")
                    api_call("DELETE", p, token=token)
    log("PASS", "Cleanup done")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    print(f"\n{'='*60}")
    print(f"  Phase 5 Demo Data Setup v2  |  BASE: {BASE_URL}")
    print(f"{'='*60}\n")

    if "--cleanup" in sys.argv:
        token = step0()
        cleanup(token)
        return

    token  = step0()
    cv_id  = step1(token)
    job_id, _ = (step2(token, cv_id) if cv_id else (None, "no_cv"))
    app_id = step3(token)
    if job_id:
        step4(token, app_id, job_id)
    step5(token, app_id)
    step6(token, app_id)
    step7(token)
    step8(token, app_id)
    step9(token, app_id)

    print(f"\n{'='*60}")
    print(f"  Demo data COMPLETE")
    print(f"  Email:    {DEMO_EMAIL}")
    print(f"  Password: {DEMO_PASS}")
    print(f"  App ID:   {app_id}")
    print(f"  CV ID:    {cv_id}")
    print(f"  Job ID:   {job_id}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
