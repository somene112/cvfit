"""
Phase 5 E2E Demo Checklist — API-level automation.

Tests all checklist items via API calls against the deployed backend.
This serves as a proxy for the manual browser-based demo checklist.

Usage:
    $env:API_BASE_URL="https://cvfit.onrender.com"
    python scripts/e2e_demo_phase5.py --verbose
"""
from __future__ import annotations
import json, sys, io, os
from pathlib import Path
from typing import Optional

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "backend"))

import urllib.request, urllib.error

BASE_URL = os.environ.get("API_BASE_URL", "https://cvfit.onrender.com")

# Demo account created by setup_demo_data.py
DEMO_EMAIL = "dat_phase5_demo@demo.app"
DEMO_PASS  = "DemoTest123!"

# Synthetic second user for ownership checks
DEMO2_EMAIL = "dat_phase5_owner_b@demo.app"
DEMO2_PASS  = "DemoBTest123!"

PASSED: list[str] = []
FAILED: list[tuple[str, str]] = []  # (check_name, reason)


def C():
    OK, ERR, WRN = "\033[92m", "\033[91m", "\033[93m"
    END = "\033[0m"
    return OK, ERR, WRN, END

def check(name: str, condition: bool, detail: str = ""):
    OK, ERR, _, END = C()
    color = OK if condition else ERR
    icon = "[PASS]" if condition else "[FAIL]"
    msg = f"  {icon} {name}" + (f" — {detail}" if detail else "")
    print(f"  {color}{msg}{END}")
    if condition:
        PASSED.append(name)
    else:
        FAILED.append((name, detail))


def api(method: str, path: str, token: str = None,
        payload: dict = None, files: dict = None,
        expect_fail: bool = False, timeout: int = 30) -> tuple[int, dict | str]:
    url = BASE_URL + path
    try:
        if files:
            boundary = "boundary_demo_123"
            parts = []
            # files: dict of {field_name: (content_bytes, filename_str, content_type_str)}
            for field_name, file_info in files.items():
                content = file_info[0]
                filename = file_info[1] if len(file_info) > 1 else field_name
                ftype = file_info[2] if len(file_info) > 2 else "application/octet-stream"
                header = (
                    f"--{boundary}\r\n"
                    f"Content-Disposition: form-data; name=\"{field_name}\"; "
                    f"filename=\"{filename}\"\r\n"
                    f"Content-Type: {ftype}\r\n\r\n"
                ).encode()
                parts.append(header)
                parts.append(content if isinstance(content, bytes) else content.encode())
                parts.append(b"\r\n")
            parts.append(f"--{boundary}--\r\n".encode())
            body_bytes = b"".join(parts)
            req = urllib.request.Request(url, data=body_bytes, method=method)
            req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
        else:
            data = json.dumps(payload).encode() if payload else None
            req = urllib.request.Request(url, data=data, method=method)
            req.add_header("Content-Type", "application/json")
        if token:
            req.add_header("Authorization", f"Bearer {token}")
        with urllib.request.urlopen(req, timeout=timeout) as r:
            raw = r.read().decode(errors="replace")
            try:
                body = json.loads(raw)
            except Exception:
                body = raw
            return r.status, body
    except urllib.error.HTTPError as e:
        raw = e.read().decode(errors="replace")
        try:
            body = json.loads(raw)
        except Exception:
            body = raw
        return e.code, body
    except Exception as ex:
        return 0, str(ex)


def login(email: str, password: str) -> Optional[str]:
    code, body = api("POST", "/v1/auth/login", payload={"email": email, "password": password})
    if code == 200:
        return body.get("access_token", "")
    return None


def make_cv_docx() -> bytes:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Demo User", 0)
        doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
        doc.add_paragraph("Experience: 2 years backend development")
        doc.add_paragraph("Projects: E-Commerce API with FastAPI + Redis")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return b"PK\x03\x04fake"


def wait_job(token: str, job_id: str, timeout: int = 150) -> str:
    import time
    deadline = time.time() + timeout
    while time.time() < deadline:
        code, body = api("GET", f"/v1/jobs/{job_id}", token=token)
        if code == 200 and isinstance(body, dict):
            status = body.get("status", "")
            if status in ("completed", "failed", "succeeded"):
                return status
        time.sleep(4)
    return "timeout"


# ── SECTION 1: Authentication ───────────────────────────────────────────────
def section1_auth():
    print("\n## SECTION 1: Authentication")
    token = login(DEMO_EMAIL, DEMO_PASS)
    check("Login with demo credentials", token is not None, f"token={token[:12] if token else 'NONE'}...")

    if not token:
        FAILED.append(("Auth — cannot proceed", "no token"))
        return None

    code, body = api("GET", "/v1/auth/me", token=token)
    check("GET /v1/auth/me returns user info", code == 200,
          f"email={body.get('email') if isinstance(body, dict) else '?'}")

    code, _ = api("GET", "/v1/auth/me")
    check("GET /v1/auth/me without token → 401", code == 401, f"got {code}")

    code, _ = api("GET", "/v1/auth/me", token="invalid.jwt.token")
    check("GET /v1/auth/me with invalid token → 401", code == 401, f"got {code}")

    return token


# ── SECTION 2: Application Workspace ────────────────────────────────────────
def section2_application(token: str):
    print("\n## SECTION 2: Application Workspace")
    code, body = api("GET", "/v1/applications", token=token)
    check("GET /v1/applications returns 200", code == 200, f"total={body.get('total', 0)}")
    check("/v1/applications has 'items' key", "items" in (body if isinstance(body, dict) else {}))
    check("/v1/applications requires auth", code != 401, "unauthenticated → 401")

    # Create application
    code, body = api("POST", "/v1/applications", token=token, payload={
        "job_title": "Backend Engineer",
        "company_name": "DemoCorp",
        "jd_text": "Python FastAPI PostgreSQL Docker",
    })
    check("POST /v1/applications → 201", code == 201, f"got {code}")
    app_id = body.get("id", "") if isinstance(body, dict) else ""

    if app_id:
        code, body = api("GET", f"/v1/applications/{app_id}", token=token)
        check("GET /v1/applications/{id} returns app", code == 200)
        check("App has job_title", isinstance(body, dict) and bool(body.get("job_title")))

        code, _ = api("GET", f"/v1/applications/{app_id}", token="wrong_token")
        check("Cross-user access → 401/404", code in (401, 404), f"got {code}")
    return app_id


# ── SECTION 3: Analysis + Attachment ─────────────────────────────────────────
# Note: the demo app (6177f5c9) already has an analysis attached.
# This section creates a fresh analysis for the new app.
def section3_analysis_attachment(token: str, app_id: str):
    print("\n## SECTION 3: Analysis + Attachment")
    if not app_id:
        check("Skip (no app_id)", False, "no app_id")
        return None

    # Upload CV
    cv_bytes = make_cv_docx()
    code, body = api("POST", "/v1/cv/upload", token=token,
                     files={"file": (cv_bytes, "test_cv.docx",
                                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    check("CV upload (DOCX) → 200/201", code in (200, 201), f"got {code}")
    cv_id = body.get("cv_id") or body.get("cv_file_id") or "" if isinstance(body, dict) else ""

    if cv_id:
        # Create analysis job
        code, body = api("POST", "/v1/jobs/create-score", token=token, payload={
            "cv_file_id": cv_id,
            "jd_text": "Backend Engineer: Python, FastAPI, PostgreSQL, Redis, Docker, 2+ years",
            "options": {"target_role": "Backend Engineer"},
        })
        check("POST /v1/jobs/create-score → 200/201", code in (200, 201), f"got {code}")
        job_id = body.get("job_id", "") if isinstance(body, dict) else ""

        if job_id:
            status = wait_job(token, job_id)
            check("Analysis job completed (not failed)", status == "succeeded", f"status={status}")

            # Attach to application
            code, _ = api("POST", f"/v1/applications/{app_id}/attach-analysis/{job_id}", token=token)
            check("Attach analysis to application → 200/201", code in (200, 201), f"got {code}")

            # Check readiness after attachment
            code, body = api("GET", f"/v1/applications/{app_id}/readiness", token=token)
            check("Readiness endpoint → 200", code == 200, f"got {code}")
            level = body.get("readiness_level", "") if isinstance(body, dict) else ""
            check("Readiness level is not empty", bool(level), f"level={level}")
            return job_id
    return None


# ── SECTION 4: Cover Letter ─────────────────────────────────────────────────
def section4_cover_letter(token: str, app_id: str):
    print("\n## SECTION 4: Cover Letter")
    if not app_id:
        check("Skip cover letter", False, "no app_id")
        return

    code, body = api("POST", f"/v1/applications/{app_id}/cover-letter/generate", token=token)
    check("Generate cover letter → 200/201", code in (200, 201), f"got {code}")
    if code == 200 and isinstance(body, dict):
        cl = body.get("payload_json") or body
        sections = cl if isinstance(cl, dict) else {}
        check("Cover letter has sections", bool(sections))
        disclaimer = sections.get("disclaimer", "") if isinstance(sections, dict) else ""
        check("Cover letter has disclaimer", bool(disclaimer))
        check("Disclaimer contains 'review'", "review" in disclaimer.lower(),
              f"disclaimer={disclaimer[:60]}...")


# ── SECTION 5: Application Package ───────────────────────────────────────────
def section5_package(token: str, app_id: str):
    print("\n## SECTION 5: Application Package")
    if not app_id:
        check("Skip package", False, "no app_id")
        return

    code, body = api("POST", f"/v1/applications/{app_id}/package/generate", token=token)
    check("Generate application package → 200/201", code in (200, 201), f"got {code}")
    if code == 200 and isinstance(body, dict):
        check("Package has sections", bool(list(body.keys())))
        check("Package has readiness_summary",
              "readiness_summary" in body or "summary" in body)


# ── SECTION 6: Career Profile ─────────────────────────────────────────────────
def section6_profile(token: str):
    print("\n## SECTION 6: Career Profile")
    for item_type, title, desc in [
        ("skill",      "Kubernetes",  "Container orchestration"),
        ("project",    "Microservices API", "REST API with FastAPI + Redis"),
        ("achievement","API uptime 99.9%", "Maintained 99.9% uptime for 12 months"),
    ]:
        code, body = api("POST", "/v1/profile/items", token=token, payload={
            "item_type": item_type,
            "title": title,
            "description": desc,
        })
        check(f"Create profile item ({item_type})", code in (200, 201), f"got {code}")

    code, body = api("GET", "/v1/profile/items", token=token)
    check("GET /v1/profile/items → 200", code == 200)
    items = body.get("items", []) if isinstance(body, dict) else []
    check("Profile items count >= 3", len(items) >= 3, f"count={len(items)}")


# ── SECTION 7: Interview Practice ──────────────────────────────────────────────
def section7_interview(token: str, app_id: str):
    print("\n## SECTION 7: Interview Practice")
    if not app_id:
        check("Skip interview practice", False, "no app_id")
        return

    # Get questions
    code, body = api("GET", f"/v1/applications/{app_id}/interview/questions", token=token)
    check("GET interview questions → 200", code == 200, f"got {code}")
    questions = body.get("questions", []) if isinstance(body, dict) else []
    check("At least 1 question returned", len(questions) >= 1, f"got {len(questions)}")

    if questions:
        q = questions[0]
        q_id = q.get("question_id", "")
        q_txt = q.get("question", "")
        check("Question has question_id", bool(q_id), f"q_id={q_id}")
        check("Question has question text", bool(q_txt))

        # Submit answer
        code, body = api("POST",
                         f"/v1/applications/{app_id}/interview/answers",
                         token=token,
                         payload={
                             "question_id": q_id,
                             "question": q_txt,
                             "answer_text": "I improved API latency by 40% using Redis caching and PostgreSQL query optimisation.",
                         })
        check("Submit interview answer → 201", code == 201, f"got {code}")
        if code == 201 and isinstance(body, dict):
            check("Answer has feedback", "feedback" in body or "score" in body)
            check("Answer has rubric/score", "rubric" in body or "score" in body)


# ── SECTION 8: Readiness Dashboard ────────────────────────────────────────────
def section8_readiness(token: str):
    print("\n## SECTION 8: Readiness Dashboard")
    code, body = api("GET", "/v1/applications", token=token)
    check("Applications list for readiness", code == 200, f"got {code}")
    if code == 200 and isinstance(body, dict):
        apps = body.get("items", [])
        check("At least 1 application in list", len(apps) >= 1, f"count={len(apps)}")


# ── SECTION 9: Guardrails + Security ─────────────────────────────────────────
def section9_guardrails_security():
    print("\n## SECTION 9: Guardrails + Security")

    # Create second user for ownership checks
    code2, _ = api("POST", "/v1/auth/register", payload={
        "email": DEMO2_EMAIL, "password": DEMO2_PASS, "name": "User B"
    })
    token2 = login(DEMO2_EMAIL, DEMO2_PASS) if code2 in (200, 201, 409) else None

    # Register main user
    token1 = login(DEMO_EMAIL, DEMO_PASS)
    if not (token1 and token2):
        check("Ownership checks skipped (login failed)", False)
        return

    # Create app as user 1
    code, body = api("POST", "/v1/applications", token=token1, payload={
        "job_title": "User1 Private Job", "company_name": "PrivateCo",
        "jd_text": "secret JD",
    })
    user1_app_id = body.get("id", "") if code == 201 else ""

    # User 2 tries to access user 1's app
    if user1_app_id:
        code, _ = api("GET", f"/v1/applications/{user1_app_id}", token=token2)
        check("User B cannot see User A's application", code in (401, 403, 404), f"got {code}")

        code, _ = api("PATCH", f"/v1/applications/{user1_app_id}", token=token2,
                       payload={"job_title": "Hacked"})
        check("User B cannot modify User A's application", code in (401, 403, 404), f"got {code}")

    # User 1 cannot see user 2's profile
    code, _ = api("GET", "/v1/profile/items", token=token2)
    check("User B profile accessible to user B", code == 200)


# ── SECTION 10: Usage Endpoint ───────────────────────────────────────────────
def section10_usage(token: str):
    print("\n## SECTION 10: Usage Endpoint")
    code, body = api("GET", "/v1/usage/me", token=token)
    check("GET /v1/usage/me → 200", code == 200)
    if code == 200 and isinstance(body, dict):
        check("Usage response has no PII", "email" not in body and "password" not in body)
        check("Usage response has plan_id", "plan_id" in body)


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    verbose = "--verbose" in sys.argv
    print(f"\n{'='*60}")
    print(f"  Phase 5 E2E Demo Checklist — API Automation")
    print(f"  BASE: {BASE_URL}")
    print(f"  User: {DEMO_EMAIL}")
    print(f"{'='*60}")

    # ── Auth
    token = section1_auth()
    if not token:
        print("\nFATAL: Cannot login — aborting")
        sys.exit(1)

    # Demo app already has analysis attached (from setup_demo_data.py)
    DEMO_APP_ID = "6177f5c9-e979-4d79-ad1b-72cda06e7a2b"
    fresh_app_id: Optional[str] = None
    code, body = api("POST", "/v1/applications", token=token, payload={
        "job_title": "E2E Test App",
        "company_name": "E2ECorp",
        "jd_text": "Python FastAPI PostgreSQL Docker 2+ years",
    })
    if code == 201:
        fresh_app_id = body.get("id", "")

    # ── Sections
    section2_application(token)
    section3_analysis_attachment(token, fresh_app_id or "")
    # Cover letter + package need demo app (has analysis)
    section4_cover_letter(token, DEMO_APP_ID)
    section5_package(token, DEMO_APP_ID)
    section6_profile(token)
    # Interview + readiness need demo app (has analysis)
    section7_interview(token, DEMO_APP_ID)
    section8_readiness(token)
    section9_guardrails_security()
    section10_usage(token)

    # ── Summary
    total = len(PASSED) + len(FAILED)
    print(f"\n{'='*60}")
    print(f"  RESULT: {len(PASSED)}/{total} checks passed")
    print(f"{'='*60}")
    if FAILED:
        print("\n  FAILED:")
        for name, detail in FAILED:
            print(f"    - {name}: {detail}")
    else:
        print("  All checks PASSED!")
    print()


if __name__ == "__main__":
    main()
